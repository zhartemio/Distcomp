"""Copy Desktop\\lab3RV.docx -> lab4RV.docx and rewrite for lab 4 (Kafka)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\ACER\Desktop\lab3RV.docx")
DST = Path(r"C:\Users\ACER\Desktop\lab4RV.docx")


def set_paragraph_text(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    for r in runs:
        r.text = ""
    if runs:
        runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_in_doc(doc: Document) -> None:
    chain_subs: list[tuple[str, str]] = [
        ("по лабораторной работе №3", "по лабораторной работе №4"),
        ("лабораторной работы №3", "лабораторной работы №4"),
        ("лабораторной работе №3", "лабораторной работе №4"),
    ]

    full_map: dict[str, str] = {
        "Тема работы: Модуляризация приложения с использованием базы данных Cassandra": (
            "Тема работы: Брокеры сообщений на примере Apache Kafka "
            "(транспорт сообщений Message между publisher и discussion)"
        ),
        "Настоящая лабораторная работа является продолжением разработки серверного приложения на языке Python для управления связанными сущностями Author, News, Mark, Message и посвящена выделению хранения сообщений Message в отдельный модуль discussion с использованием распределённой СУБД Apache Cassandra, при сохранении основной части системы publisher на реляционной СУБД PostgreSQL.": (
            "Настоящая лабораторная работа развивает решение лабораторной работы №3 и посвящена замене прямого HTTP-обмена между модулями publisher и discussion на асинхронный транспорт через брокер Apache Kafka. "
            "Сущности Author, News, Mark по-прежнему обслуживаются модулем publisher (PostgreSQL), а Message хранится в Cassandra в модуле discussion; при этом операции с сообщениями для внешнего клиента по-прежнему доступны через REST publisher."
        ),
        "Цель работы — реализовать микросервисную архитектуру: основное приложение продолжает обслуживать сущности Author, News, Mark и связи News–Mark в PostgreSQL, а сущность Message переносится в сервис discussion, где данные хранятся в Cassandra с учётом ключа партиционирования и кластеризации, заданных в условии лабораторной работы.": (
            "Цель работы — внедрить Kafka-топики InTopic и OutTopic для обмена командами и результатами по сущности Message между publisher и discussion, сохранить внешний REST-контракт и порты 24110/24130, "
            "добавить для сообщения поле состояния state (в т.ч. результаты автоматической модерации контента), обеспечить согласованное партиционирование записей по ключу, связанному с newsId, и подготовить docker-compose стек с Zookeeper и Kafka."
        ),
        "Обеспечивается взаимодействие модулей по REST, сохранение внешнего API с префиксом /api/v1.0/ на стороне publisher на localhost:24110 и отдельный REST-интерфейс discussion на localhost:24130, а также инициализация схемы Cassandra с keyspace distcomp и таблицами с префиксом tbl_.": (
            "Внешний клиент и тесты обращаются к publisher по REST с префиксом /api/v1.0/ (localhost:24110); модуль discussion сохраняет собственный REST на localhost:24130. "
            "Внутренняя доставка операций Message между publisher и discussion выполняется через Kafka: запросы публикуются в InTopic, ответы читаются из OutTopic по correlationId с таймаутом. "
            "Схема Cassandra (keyspace distcomp, таблицы tbl_) дополнена хранением state для сообщений."
        ),
        "В ходе выполнения работы были решены следующие задачи: разделение решения на два сервиса publisher и discussion; проектирование модели Message для Cassandra; реализация REST-API в discussion и проксирование операций из publisher через HTTP-клиент httpx; настройка Docker Compose; обеспечение обратной совместимости внешних маршрутов publisher для сообщений; проверка работоспособности автоматизированными тестами учебного тренажёра.": (
            "В ходе выполнения работы были решены следующие задачи: настройка окружения Apache Kafka (образы Confluent, сеть Docker); "
            "реализация продюсера/консьюмера на publisher (отправка в InTopic, ожидание ответа из OutTopic); реализация фонового обработчика на discussion (чтение InTopic, выполнение операций через общий сервисный слой с REST, ответ в OutTopic); "
            "модерация текста сообщения и запись state; ключ партиции Kafka по newsId; расширение docker-compose (zookeeper, kafka, healthcheck); сохранение совместимости JSON-ответов с автотестами; проверка сценариев тренажёра."
        ),
        "1.2 Проектирование микросервисного взаимодействия\t5": "1.2 Проектирование обмена через Kafka\t5",
        "1.3 Проектирование схемы Cassandra\t5": "1.3 Схема Cassandra, поле state и модерация\t5",
        "1.4 Инициализация баз данных\t6": "1.4 Инициализация баз данных и топиков\t6",
        "2.2 Структура программного продукта\t7": "2.2 Структура программного продукта (каталог task4)\t7",
        "2.3 Реализация REST и совместимости API\t7": "2.3 REST, Kafka-RPC и совместимость ответов\t7",
        "Взаимодействие организовано по принципу «publisher как фасад для клиента»: клиент и автоматические тесты обращаются к publisher по привычным URL сообщений.": (
            "Взаимодействие с точки зрения клиента остаётся прежним: publisher выступает фасадом, тесты обращаются к тем же URL сообщений на порту 24110."
        ),
        "При обработке запросов к сущности Message publisher проверяет наличие связанной News в PostgreSQL и затем перенаправляет операции CRUD к discussion по внутреннему адресу, задаваемому через переменную окружения DISCUSSION_BASE_URL.": (
            "При обработке запросов к Message publisher проверяет наличие связанной News в PostgreSQL и формирует конверт {correlationId, op, payload}, который отправляется в топик InTopic с ключом партиции, привязанным к newsId (где применимо). "
            "Ответ ожидается из топика OutTopic по совпадению correlationId. Для отладки без брокера предусмотрен опциональный режим прямого HTTP к discussion через переменную DISCUSSION_BASE_URL."
        ),
        "Обработка сбоев реализована через преобразование ответов discussion в HTTPException с телом errorMessage и errorCode, согласованным с остальным API.": (
            "Ошибки и таймауты обмена через Kafka преобразуются в HTTPException с полями errorMessage и errorCode в том же стиле, что и остальной API publisher."
        ),
        "1.3 Проектирование схемы Cassandra": "1.3 Схема Cassandra, поле state и модерация",
        "Таблица tbl_message в keyspace distcomp описывается первичным ключом ((country), news_id, id), что соответствует условию задания и обеспечивает логичное размещение данных по партициям.": (
            "Таблица tbl_message в keyspace distcomp сохраняет первичный ключ ((country), news_id, id). Дополнительно введена колонка state для хранения результата модерации (APPROVE/DECLINE); при обновлении существующей установки выполняется ALTER TABLE при необходимости."
        ),
        "Для соответствия сценариям REST и тестовым проверкам в реализации применяются правила выдачи следующего идентификатора сообщения и однозначного чтения по id с учётом возможных коллизий.": (
            "Автоматическая модерация реализована отдельным модулем (например, отклонение при наличии явного служебного маркера в тексте). "
            "Сохранены правила нумерации id и однозначного чтения по id, использованные в лабораторной работе №3."
        ),
        "1.4 Инициализация баз данных": "1.4 Инициализация баз данных и топиков",
        "Для Cassandra при первом обращении к хранилищу выполняется создание keyspace и таблицы tbl_message, если они ещё не существуют.": (
            "Для Cassandra при первом обращении выполняется создание keyspace и таблицы tbl_message при необходимости. Топики InTopic и OutTopic создаются при первой публикации (автосоздание на стороне брокера)."
        ),
        "Для NoSQL-хранилища использованы Apache Cassandra 4.x и драйвер cassandra-driver, для межсервисных вызовов — httpx, для контейнеризации — Docker и Docker Compose.": (
            "Для NoSQL-хранилища использованы Apache Cassandra 4.x и драйвер cassandra-driver; для обмена между publisher и discussion — библиотека kafka-python и брокер Apache Kafka; httpx оставлен в publisher для опционального HTTP-fallback. Контейнеризация — Docker и Docker Compose."
        ),
        "Каталог task3 содержит два приложения: publisher и discussion.": (
            "Каталог task4 содержит два приложения: publisher и discussion и файл docker-compose.yml для полного стека (Postgres, Cassandra, Zookeeper, Kafka, publisher, discussion)."
        ),
        "Файл docker-compose.yml оркестрирует сервисы postgres, cassandra, publisher и discussion. Архитектура соответствует многослойному подходу: контроллеры принимают и отдают DTO, сервисы содержат бизнес-правила, а репозитории и низкоуровневые модули инкапсулируют работу с данными.": (
            "Файл docker-compose.yml оркестрирует сервисы postgres, cassandra, zookeeper, kafka, publisher и discussion с общей сетью и healthcheck для устойчивого старта. Архитектура модулей по слоям сохранена; в discussion добавлен фоновый потребитель InTopic."
        ),
        "В publisher сохранены маршруты сообщений на localhost:24110, но реализация делегирована discussion; при необходимости в тело запроса добавляется или прокидывается поле country через значение по умолчанию из переменной окружения.": (
            "В publisher сохранены маршруты сообщений на localhost:24110; делегирование в discussion выполняется через Kafka-RPC (или через HTTP при заданном DISCUSSION_BASE_URL). Поле country по умолчанию задаётся на стороне discussion."
        ),
        "docker-compose.yml описывает запуск контейнеров PostgreSQL и Cassandra, а также сборку образов publisher и discussion, проброс портов 5432, 9042, 24110 и 24130.": (
            "docker-compose.yml описывает запуск PostgreSQL, Cassandra, Zookeeper, Kafka, а также сборку образов publisher и discussion; пробрасываются порты 5432, 9042, 9092, 2181, 24110 и 24130."
        ),
        "Через переменные окружения задаются строки подключения DATABASE_URL, базовый URL DISCUSSION_BASE_URL, параметры CASSANDRA_HOSTS и значение MESSAGE_DEFAULT_COUNTRY.": (
            "Через переменные окружения задаются DATABASE_URL, CASSANDRA_HOSTS, MESSAGE_DEFAULT_COUNTRY, KAFKA_BOOTSTRAP_SERVERS, KAFKA_IN_TOPIC, KAFKA_OUT_TOPIC, KAFKA_RPC_TIMEOUT_SEC, опционально DISCUSSION_BASE_URL и идентификаторы групп consumer."
        ),
        "Тестирование выполнялось с помощью Docker Compose: поднимались PostgreSQL с базой distcomp, Cassandra, сервис publisher на порту 24110 и discussion на порту 24130.": (
            "Тестирование выполнялось с помощью Docker Compose: поднимались PostgreSQL (distcomp), Cassandra, Zookeeper, Kafka, publisher (24110) и discussion (24130)."
        ),
        "Автоматические тесты учебного тренажёра выполняли HTTP-запросы к API publisher, а также прямые проверки к Cassandra, включая паузы ожидания готовности кластера.": (
            "Автоматические тесты тренажёра выполняли HTTP-запросы к API publisher; при необходимости выполнялись проверки данных в Cassandra. Учитывалась готовность Kafka и фонового consumer в discussion."
        ),
        "Для стабильности старта были настроены healthcheck для Cassandra и проверка готовности discussion перед запуском publisher.": (
            "Для стабильности старта настроены healthcheck для Cassandra и Kafka, зависимость publisher от healthy discussion и kafka, увеличенный таймаут RPC и повторное подключение consumer при сбоях сессии."
        ),
        "В процессе отладки были выявлены и устранены типовые проблемы интеграции: недоступность discussion при раннем старте publisher, ошибки при обращении к строкам результата Cassandra и некорректная интерпретация id сообщения при наличии нескольких строк с одинаковым id в разных news_id.": (
            "В процессе отладки учитывались задержки готовности брокера и consumer, порядок seek/partition assignment на стороне publisher при чтении OutTopic, а также ранее выявленные в лабораторной работе №3 особенности Cassandra и идентификаторов сообщений."
        ),
        "Для устранения проблем были настроены порядок зависимостей и healthcheck в Compose, реализованы ленивое подключение к Cassandra и увеличение времени ожидания готовности базы данных, а также введены глобальная нумерация id в рамках страны и детерминированный выбор строки при выборке по id.": (
            "Для устранения проблем скорректированы depends_on и healthcheck, увеличены таймауты ожидания ответа из OutTopic, добавлен цикл перезапуска Kafka-сессии в discussion, для JSON-ответов применена политика сериализации поля state, совместимая с автотестами."
        ),
        "Ошибки проксирования publisher → discussion были обёрнуты в обработку сетевых сбоев с понятным JSON-ответом. После исправлений все тестовые сценарии по сообщениям и остальным сущностям успешно прошли.": (
            "Ошибки транспорта (в т.ч. таймаут Kafka-RPC) возвращаются клиенту в согласованном формате. После настройки стека тестовые сценарии по сообщениям и остальным сущностям успешно проходят (результаты приведены на рисунке 1 — вставить скриншот из тренажёра)."
        ),
        "В ходе выполнения лабораторной работы №3 выполнен переход от монолитного хранения сообщений в PostgreSQL к модульной архитектуре: реализованы сервисы publisher для работы с Author, News, Mark и связью News–Mark в PostgreSQL и discussion для работы с Message в Cassandra.": (
            "В ходе выполнения лабораторной работы №4 на базе модульной архитектуры лабораторной работы №3 внедрён брокер Apache Kafka для транспорта операций Message между publisher и discussion при сохранении PostgreSQL и Cassandra."
        ),
        "Настроено REST-взаимодействие между сервисами, сохранён внешний контракт API на /api/v1.0/ и портах 24110 / 24130, а также подготовлена контейнерная среда запуска.": (
            "Сохранены внешние REST-интерфейсы и контракт /api/v1.0/ на портах 24110 и 24130; расширена контейнерная среда за счёт Zookeeper и Kafka."
        ),
        "Спроектирована и реализована распределённая модель хранения сообщений в Cassandra с ключом ((country), news_id, id). Сохранена и доработана реляционная часть в PostgreSQL с префиксом таблиц tbl_ и миграциями.": (
            "Модель хранения Message в Cassandra с ключом ((country), news_id, id) сохранена; добавлено поле state и модерация. Реляционная часть в PostgreSQL с префиксом tbl_ и миграциями Alembic сохранена."
        ),
        "Реализован прозрачный для клиента вызов операций с сообщениями через publisher с делегированием в discussion, настроены Docker и Docker Compose для воспроизводимого поднятия всего стека, проведена отладка по результатам автоматических тестов.": (
            "Реализован прозрачный для клиента вызов операций с сообщениями через publisher с делегированием в discussion по Kafka, настроен docker-compose для воспроизводимого поднятия стека, выполнена проверка тестами тренажёра."
        ),
        "Таким образом, поставленные цели лабораторной работы №3 достигнуты: приложение соответствует требованиям по модуляризации и использованию Cassandra для сущности Message, при этом основная часть доменной модели остаётся в PostgreSQL.": (
            "Таким образом, поставленные цели лабораторной работы №4 достигнуты: для сущности Message обеспечен транспорт через топики InTopic/OutTopic Apache Kafka при сохранении REST, Cassandra и PostgreSQL в соответствии с заданием."
        ),
        "По условию лабораторной работы №3 сущность Message должна храниться в Cassandra в отдельном модуле discussion, тогда как Author, News, Mark и таблица связи остаются в PostgreSQL в модуле publisher.": (
            "По условию лабораторной работы №4 транспорт Message между publisher и discussion выполняется через Kafka; сущность Message по-прежнему хранится в Cassandra в discussion, а Author, News, Mark и связь News–Mark — в PostgreSQL в publisher."
        ),
    }

    def process_paragraph(p) -> None:
        raw = p.text
        t = raw.replace("\u00a0", " ").strip()
        if t in full_map:
            set_paragraph_text(p, full_map[t])
            return
        new_t = raw
        for old, new in chain_subs:
            new_t = new_t.replace(old, new)
        if new_t != raw:
            set_paragraph_text(p, new_t)

    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                process_paragraph(p)


def main() -> None:
    if not SRC.is_file():
        raise SystemExit(f"Not found: {SRC}")
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    replace_in_doc(doc)
    doc.save(DST)
    print(f"Written: {DST}")


if __name__ == "__main__":
    main()
