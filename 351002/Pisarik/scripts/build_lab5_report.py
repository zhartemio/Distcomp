"""Copy Desktop\\lab4RV.docx (or lab3RV) -> lab5RV.docx and rewrite for lab 5 (Redis cache)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

DESKTOP = Path.home() / "Desktop"
SRC = DESKTOP / "lab4RV.docx"
if not SRC.is_file():
    SRC = DESKTOP / "lab3RV.docx"
DST = DESKTOP / "lab5RV.docx"


def set_paragraph_text(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    for r in runs:
        r.text = ""
    if runs:
        runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_in_doc(doc: Document) -> None:
    chain_subs: list[tuple[str, str]] = [
        ("по лабораторной работе №4", "по лабораторной работе №5"),
        ("лабораторной работы №4", "лабораторной работы №5"),
        ("лабораторной работе №4", "лабораторной работе №5"),
        ("по лабораторной работе №3", "по лабораторной работе №5"),
        ("лабораторной работы №3", "лабораторной работы №5"),
        ("лабораторной работе №3", "лабораторной работе №5"),
        ("каталог task4", "каталог task5"),
        ("task4", "task5"),
    ]

    full_map: dict[str, str] = {
        "Тема работы: Брокеры сообщений на примере Apache Kafka (транспорт сообщений Message между publisher и discussion)": (
            "Тема работы: Кеширование данных на примере Redis в распределённом приложении "
            "(publisher: PostgreSQL + Redis; Message: Kafka ↔ Cassandra)"
        ),
        "Тема работы: Модуляризация приложения с использованием базы данных Cassandra": (
            "Тема работы: Кеширование данных на примере Redis в распределённом приложении "
            "(publisher: PostgreSQL + Redis; Message: Kafka ↔ Cassandra)"
        ),
        "Настоящая лабораторная работа развивает решение лабораторной работы №3 и посвящена замене прямого HTTP-обмена между модулями publisher и discussion на асинхронный транспорт через брокер Apache Kafka. Сущности Author, News, Mark по-прежнему обслуживаются модулем publisher (PostgreSQL), а Message хранится в Cassandra в модуле discussion; при этом операции с сообщениями для внешнего клиента по-прежнему доступны через REST publisher.": (
            "Настоящая лабораторная работа развивает решение лабораторной работы №4 и посвящена внедрению кеша Redis на стороне модуля publisher в соответствии с диаграммой взаимодействия Task 350. "
            "Сохранены REST-интерфейсы /api/v1.0/, порты 24110 и 24130, транспорт сообщений Message через Kafka между publisher и discussion и хранение в Cassandra; поверх этого реализовано кеширование ответов для сущностей из PostgreSQL и результатов операций с сообщениями после обмена через Kafka."
        ),
        "Цель работы — внедрить Kafka-топики InTopic и OutTopic для обмена командами и результатами по сущности Message между publisher и discussion, сохранить внешний REST-контракт и порты 24110/24130, добавить для сообщения поле состояния state (в т.ч. результаты автоматической модерации контента), обеспечить согласованное партиционирование записей по ключу, связанному с newsId, и подготовить docker-compose стек с Zookeeper и Kafka.": (
            "Цель работы — реализовать на publisher стратегию cache-aside с использованием Redis: чтение из кеша перед обращением к PostgreSQL для сущностей Author, News, Mark и инвалидация/обновление кеша при изменениях; "
            "для Message — обращение к Redis перед отправкой запроса в Kafka, сохранение успешного ответа из OutTopic в Redis, согласованная инвалидация ключей при CRUD; развернуть Redis официальным образом с Docker Hub и расширить docker-compose."
        ),
        "2.2 Структура программного продукта (каталог task5)\t7": "2.2 Структура программного продукта (каталог task5)\t7",
        "2.3 REST, Kafka-RPC и совместимость ответов\t7": "2.3 REST, Kafka, Redis и согласованность кеша\t7",
        "2.4 Конфигурация и окружение\t8": "2.4 Конфигурация и окружение (REDIS_URL)\t8",
        "docker-compose.yml описывает запуск PostgreSQL, Cassandra, Zookeeper, Kafka, а также сборку образов publisher и discussion; пробрасываются порты 5432, 9042, 9092, 2181, 24110 и 24130.": (
            "docker-compose.yml описывает запуск PostgreSQL, Cassandra, Zookeeper, Kafka, Redis (образ redis:7-alpine, порт 6379), а также сборку образов publisher и discussion; пробрасываются порты 5432, 9042, 9092, 2181, 6379, 24110 и 24130."
        ),
        "Через переменные окружения задаются DATABASE_URL, CASSANDRA_HOSTS, MESSAGE_DEFAULT_COUNTRY, KAFKA_BOOTSTRAP_SERVERS, KAFKA_IN_TOPIC, KAFKA_OUT_TOPIC, KAFKA_RPC_TIMEOUT_SEC, опционально DISCUSSION_BASE_URL и идентификаторы групп consumer.": (
            "Через переменные окружения задаются DATABASE_URL, REDIS_URL (по умолчанию redis://redis:6379/0), параметры Kafka и Cassandra, MESSAGE_DEFAULT_COUNTRY, опционально DISCUSSION_BASE_URL."
        ),
        "Тестирование выполнялось с помощью Docker Compose: поднимались PostgreSQL (distcomp), Cassandra, Zookeeper, Kafka, publisher (24110) и discussion (24130).": (
            "Тестирование выполнялось с помощью Docker Compose: поднимались PostgreSQL, Cassandra, Zookeeper, Kafka, Redis, publisher и discussion."
        ),
        "Автоматические тесты тренажёра выполняли HTTP-запросы к API publisher; при необходимости выполнялись проверки данных в Cassandra. Учитывалась готовность Kafka и фонового consumer в discussion.": (
            "Автоматические тесты тренажёра выполняли HTTP-запросы к API publisher, в том числе сценарий проверки цепочки REST + Redis + Kafka + Cassandra для сущности Message; учитывались healthcheck Redis и порядок инвалидации кеша."
        ),
        "Для стабильности старта настроены healthcheck для Cassandra и Kafka, зависимость publisher от healthy discussion и kafka, увеличенный таймаут RPC и повторное подключение consumer при сбоях сессии.": (
            "Для стабильности старта настроены healthcheck для Cassandra, Kafka и Redis; publisher ожидает готовность зависимостей перед стартом; сохранены таймауты Kafka-RPC и устойчивость consumer в discussion."
        ),
        "Таким образом, поставленные цели лабораторной работы №4 достигнуты: для сущности Message обеспечен транспорт через топики InTopic/OutTopic Apache Kafka при сохранении REST, Cassandra и PostgreSQL в соответствии с заданием.": (
            "Таким образом, поставленные цели лабораторной работы №5 достигнуты: в модуле publisher внедрён Redis-кеш для ускорения повторяемых запросов и согласованного обновления данных при операциях с Message через Kafka и Cassandra."
        ),
    }

    def process_paragraph(p) -> None:
        raw = p.text
        t = raw.replace("\u00a0", " ").strip()
        if t in full_map:
            set_paragraph_text(p, full_map[t])
            return
        new_t = raw
        for old, new in chain_subs:
            new_t = new_t.replace(old, new)
        if new_t != raw:
            set_paragraph_text(p, new_t)

    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                process_paragraph(p)


def main() -> None:
    if not SRC.is_file():
        raise SystemExit(f"Source not found: {SRC} (нужен lab4RV.docx или lab3RV.docx на рабочем столе)")
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    replace_in_doc(doc)
    doc.save(DST)
    print(f"Written: {DST} (from {SRC.name})")


if __name__ == "__main__":
    main()
