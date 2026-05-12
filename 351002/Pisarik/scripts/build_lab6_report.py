"""Copy Desktop\\lab5RV.docx -> lab6RV.docx; rewrite for lab 6 (Security / JWT)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

DESKTOP = Path.home() / "Desktop"
SRC = DESKTOP / "lab5RV.docx"
if not SRC.is_file():
    SRC = DESKTOP / "lab4RV.docx"
if not SRC.is_file():
    SRC = DESKTOP / "lab3RV.docx"
DST = DESKTOP / "lab6RV.docx"


def set_paragraph_text(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    for r in runs:
        r.text = ""
    if runs:
        runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_in_doc(doc: Document) -> None:
    chain = [
        ("по лабораторной работе №5", "по лабораторной работе №6"),
        ("лабораторной работы №5", "лабораторной работы №6"),
        ("лабораторной работе №5", "лабораторной работе №6"),
        ("по лабораторной работе №4", "по лабораторной работе №6"),
        ("лабораторной работы №4", "лабораторной работы №6"),
        ("лабораторной работе №4", "лабораторной работе №6"),
        ("каталог task5", "каталог task6"),
        ("task5", "task6"),
        ("task4", "task6"),
    ]

    full_map: dict[str, str] = {
        "Тема работы: Кеширование данных на примере Redis в распределённом приложении (publisher: PostgreSQL + Redis; Message: Kafka ↔ Cassandra)": (
            "Тема работы: Интеграция поддержки Security в REST API (JWT, роли ADMIN/CUSTOMER, "
            "версия /api/v2.0 на порту 24110; /api/v1.0 без изменений)"
        ),
        "Таким образом, поставленные цели лабораторной работы №5 достигнуты: в модуле publisher внедрён Redis-кеш для ускорения повторяемых запросов и согласованного обновления данных при операциях с Message через Kafka и Cassandra.": (
            "Таким образом, поставленные цели лабораторной работы №6 достигнуты: реализованы регистрация и вход по JWT, "
            "защита эндпоинтов /api/v2.0, хранение паролей в BCrypt, разграничение прав ADMIN и CUSTOMER в соответствии с Task 361."
        ),
    }

    def proc(p) -> None:
        raw = p.text
        t = raw.replace("\u00a0", " ").strip()
        if t in full_map:
            set_paragraph_text(p, full_map[t])
            return
        new_t = raw
        for a, b in chain:
            new_t = new_t.replace(a, b)
        if new_t != raw:
            set_paragraph_text(p, new_t)

    for p in doc.paragraphs:
        proc(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    proc(p)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                proc(p)


def main() -> None:
    if not SRC.is_file():
        raise SystemExit(f"Нет исходного отчёта на рабочем столе: lab5RV / lab4RV / lab3RV.docx")
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    replace_in_doc(doc)
    doc.save(DST)
    print(f"Written: {DST} (from {SRC.name})")


if __name__ == "__main__":
    main()
