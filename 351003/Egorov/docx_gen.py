import os
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_all_py_files_content():
    py_files = glob.glob("*.py")
    all_code = ""

    for py_file in py_files:
        with open(py_file, 'r', encoding='utf-8') as file:
            content = file.read()
            all_code += f"\nФайл: {py_file}\n\n"
            all_code += content
            all_code += "\n\n"

    return all_code


def create_document_with_code():
    if not os.path.exists("base.docx"):
        print("Ошибка: файл base.docx не найден в текущей папке!")
        return

    doc = Document("base.docx")

    number = input("Введите цифру для вставки после символа №: ")

    for paragraph in doc.paragraphs:
        if '№' in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if '№' in inline[i].text:
                    text = inline[i].text
                    inline[i].text = text.replace('№', f'№{number}')
                    break

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = title.add_run("Код программы:")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    code_content = get_all_py_files_content()

    if code_content.strip():
        code_paragraph = doc.add_paragraph()
        run = code_paragraph.add_run(code_content)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    else:
        warning = doc.add_paragraph()
        run = warning.add_run("В текущей папке не найдены .py файлы")
        run.italic = True
        run.font.size = Pt(10)

    output_filename = "result.docx"
    doc.save(output_filename)
    print(f"Документ успешно создан: {output_filename}")

    py_files = glob.glob("*.py")
    if py_files:
        print(f"Найдены .py файлы: {', '.join(py_files)}")
    else:
        print("Внимание: .py файлы не найдены в текущей папке")


if __name__ == "__main__":
    create_document_with_code()